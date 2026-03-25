import os
from abc import ABC, abstractmethod
from typing import Dict, Any, Optional, List
import warnings
warnings.filterwarnings('ignore')

try:
    import PyPDF2
except ImportError:
    PyPDF2 = None

try:
    from docx import Document
except ImportError:
    Document = None

try:
    from striprtf.striprtf import rtf_to_text
except ImportError:
    rtf_to_text = None

try:
    from bs4 import BeautifulSoup
except ImportError:
    BeautifulSoup = None


class DocumentParseError(Exception):
    """Исключение при парсинге документа"""
    pass


class FileReaderStrategy(ABC):
    """Абстрактная стратегия чтения файлов"""
    
    @abstractmethod
    def read(self, file_path: str) -> str:
        """Прочитать текст из файла"""
        pass
    
    @abstractmethod
    def get_supported_extensions(self) -> List[str]:
        """Вернуть список поддерживаемых расширений"""
        pass
    
    def validate(self, file_path: str) -> bool:
        """Проверить, можно ли прочитать файл"""
        if not os.path.exists(file_path):
            return False
        ext = os.path.splitext(file_path)[1].lower()
        return ext in self.get_supported_extensions()


class TXTReaderStrategy(FileReaderStrategy):
    """Стратегия чтения TXT файлов"""
    
    def get_supported_extensions(self) -> List[str]:
        return ['.txt', '.text']
    
    def read(self, file_path: str) -> str:
        """Чтение TXT с автоматическим определением кодировки"""
        encodings = ['utf-8', 'cp1251', 'windows-1251', 'koi8-r', 'iso-8859-5', 'utf-16']
        
        for encoding in encodings:
            try:
                with open(file_path, 'r', encoding=encoding) as f:
                    return f.read()
            except (UnicodeDecodeError, UnicodeError):
                continue
        
        with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
            return f.read()


class RTFReaderStrategy(FileReaderStrategy):
    """Стратегия чтения RTF файлов"""
    
    def get_supported_extensions(self) -> List[str]:
        return ['.rtf']
    
    def read(self, file_path: str) -> str:
        if rtf_to_text is None:
            raise DocumentParseError("striprtf library not installed. Install with: pip install striprtf")
        
        encodings = ['utf-8', 'cp1251', 'windows-1251', 'koi8-r', 'utf-16']
        
        with open(file_path, 'rb') as f:
            rtf_bytes = f.read()
        
        for encoding in encodings:
            try:
                rtf_text = rtf_bytes.decode(encoding)
                plain_text = rtf_to_text(rtf_text)
                if plain_text and plain_text.strip():
                    return plain_text
            except (UnicodeDecodeError, Exception):
                continue
        
        rtf_text = rtf_bytes.decode('utf-8', errors='ignore')
        return rtf_to_text(rtf_text)


class PDFReaderStrategy(FileReaderStrategy):
    """Стратегия чтения PDF файлов"""
    
    def get_supported_extensions(self) -> List[str]:
        return ['.pdf']
    
    def read(self, file_path: str) -> str:
        if PyPDF2 is None:
            raise DocumentParseError("PyPDF2 library not installed. Install with: pip install PyPDF2")
        
        text_parts = []
        
        try:
            with open(file_path, 'rb') as f:
                pdf_reader = PyPDF2.PdfReader(f)
                
                for page_num, page in enumerate(pdf_reader.pages):
                    try:
                        page_text = page.extract_text()
                        if page_text and page_text.strip():
                            text_parts.append(f"[Страница {page_num + 1}]\n{page_text}")
                    except Exception as e:
                        print(f"Error extracting page {page_num}: {e}")
                        continue
                
                return '\n\n'.join(text_parts)
                
        except Exception as e:
            raise DocumentParseError(f"Error reading PDF: {e}")


class DOCXReaderStrategy(FileReaderStrategy):
    """Стратегия чтения DOCX файлов"""
    
    def get_supported_extensions(self) -> List[str]:
        return ['.docx']
    
    def read(self, file_path: str) -> str:
        if Document is None:
            raise DocumentParseError("python-docx library not installed. Install with: pip install python-docx")
        
        try:
            doc = Document(file_path)
            text_parts = []
            
            for para in doc.paragraphs:
                if para.text.strip():
                    text_parts.append(para.text)
            
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        if cell.text.strip():
                            row_text.append(cell.text.strip())
                    if row_text:
                        text_parts.append(' | '.join(row_text))
            
            return '\n\n'.join(text_parts)
            
        except Exception as e:
            raise DocumentParseError(f"Error reading DOCX: {e}")


class DOCReaderStrategy(FileReaderStrategy):
    """Стратегия чтения DOC файлов (старый формат)"""
    
    def get_supported_extensions(self) -> List[str]:
        return ['.doc']
    
    def read(self, file_path: str) -> str:
        try:
            import subprocess
            result = subprocess.run(
                ['antiword', file_path],
                capture_output=True,
                text=True,
                timeout=30
            )
            if result.returncode == 0 and result.stdout.strip():
                return result.stdout
        except (subprocess.SubprocessError, FileNotFoundError):
            pass
        
        try:
            with open(file_path, 'rb') as f:
                content = f.read()
                text = content.decode('utf-8', errors='ignore')
                text = ''.join(c for c in text if c.isprintable() or c in '\n\r\t')
                return text
        except Exception as e:
            raise DocumentParseError(f"Error reading DOC: {e}")


class HTMLReaderStrategy(FileReaderStrategy):
    """Стратегия чтения HTML файлов"""
    
    def get_supported_extensions(self) -> List[str]:
        return ['.html', '.htm']
    
    def read(self, file_path: str) -> str:
        if BeautifulSoup is None:
            raise DocumentParseError("beautifulsoup4 library not installed. Install with: pip install beautifulsoup4")
        
        try:
            encodings = ['utf-8', 'cp1251', 'windows-1251']
            content = None
            
            for encoding in encodings:
                try:
                    with open(file_path, 'r', encoding=encoding) as f:
                        content = f.read()
                        break
                except UnicodeDecodeError:
                    continue
            
            if content is None:
                with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                    content = f.read()
            
            soup = BeautifulSoup(content, 'html.parser')
            
            for script in soup(["script", "style"]):
                script.decompose()
            
            text = soup.get_text()
            
            lines = (line.strip() for line in text.splitlines())
            chunks = (phrase.strip() for line in lines for phrase in line.split("  "))
            text = '\n'.join(chunk for chunk in chunks if chunk)
            
            return text
            
        except Exception as e:
            raise DocumentParseError(f"Error reading HTML: {e}")


class DocumentParser:
    """
    Основной класс для парсинга документов с использованием паттерна Стратегия.
    Поддерживает динамическую регистрацию новых форматов.
    """
    
    def __init__(self):
        self._strategies: Dict[str, FileReaderStrategy] = {}
        self._register_default_strategies()
    
    def _register_default_strategies(self):
        """Регистрация стандартных стратегий"""
        strategies = [
            TXTReaderStrategy(),
            RTFReaderStrategy(),
            PDFReaderStrategy(),
            DOCXReaderStrategy(),
            DOCReaderStrategy(),
            HTMLReaderStrategy()
        ]
        
        for strategy in strategies:
            self.register_strategy(strategy)
    
    def register_strategy(self, strategy: FileReaderStrategy):
        """Зарегистрировать новую стратегию"""
        for ext in strategy.get_supported_extensions():
            self._strategies[ext] = strategy
    
    def get_strategy(self, file_path: str) -> Optional[FileReaderStrategy]:
        """Получить стратегию для файла"""
        ext = os.path.splitext(file_path)[1].lower()
        return self._strategies.get(ext)
    
    def parse(self, file_path: str, metadata: Optional[Dict] = None) -> Dict[str, Any]:
        """
        Парсит документ и возвращает текст вместе с метаданными
        
        Args:
            file_path: путь к файлу
            metadata: дополнительные метаданные
            
        Returns:
            Dict с ключами:
                - text: извлеченный текст
                - metadata: метаданные документа
                - format: формат файла
                - success: успешность операции
                - error: сообщение об ошибке (если есть)
        """
        result = {
            'text': '',
            'metadata': metadata or {},
            'format': os.path.splitext(file_path)[1].lower(),
            'success': False,
            'error': None,
            'file_size': os.path.getsize(file_path) if os.path.exists(file_path) else 0
        }
        
        if not os.path.exists(file_path):
            result['error'] = f"File not found: {file_path}"
            return result
        
        strategy = self.get_strategy(file_path)
        if not strategy:
            result['error'] = f"No strategy found for format: {result['format']}"
            return result
        
        try:
            text = strategy.read(file_path)
            
            if not text or not text.strip():
                result['error'] = "Extracted text is empty"
                return result
            
            result['text'] = text
            result['success'] = True
            result['metadata'].update({
                'format': result['format'],
                'char_count': len(text),
                'word_count': len(text.split()),
                'line_count': len(text.splitlines())
            })
            
        except DocumentParseError as e:
            result['error'] = str(e)
        except Exception as e:
            result['error'] = f"Unexpected error: {e}"
        
        return result
    
    def parse_batch(self, file_paths: List[str], metadata_list: Optional[List[Dict]] = None) -> List[Dict]:
        """Парсит несколько документов"""
        results = []
        for i, file_path in enumerate(file_paths):
            metadata = metadata_list[i] if metadata_list and i < len(metadata_list) else None
            results.append(self.parse(file_path, metadata))
        return results
    
    def get_supported_formats(self) -> List[str]:
        """Получить список поддерживаемых форматов"""
        return list(self._strategies.keys())
    
    def format_supported(self, file_path: str) -> bool:
        """Проверить, поддерживается ли формат файла"""
        return self.get_strategy(file_path) is not None


default_parser = DocumentParser()


def parse_document(file_path: str, metadata: Optional[Dict] = None) -> Dict[str, Any]:
    """
    Удобная функция для быстрого парсинга документа
    
    Example:
        result = parse_document('document.pdf')
        if result['success']:
            text = result['text']
            print(f"Extracted {result['metadata']['word_count']} words")
    """
    return default_parser.parse(file_path, metadata)